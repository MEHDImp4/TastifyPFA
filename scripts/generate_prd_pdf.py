"""
Generates a professional PDF from the TastifyPFA PRD markdown file.
Uses python-docx to create a styled Word document, then exports to PDF.
"""
import os
import sys
from pathlib import Path

import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

# Add project root to path for imports
PROJECT_ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(PROJECT_ROOT))

from markdown.extensions import Extension
from markdown.blockprocessors import BlockProcessor
from html import escape

# Configuration
PRD_FILE = PROJECT_ROOT / "docs" / "PRD_TastifyPFA.md"
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "PRD_TastifyPFA.docx"


def create_styled_document():
    """Create a professional Word document for the PRD."""
    doc = Document()

    # Define colors
    PRIMARY = RGBColor(0x23, 0x7A, 0x9E)
    SECONDARY = RGBColor(0x2D, 0x6A, 0x4F)
    DARK = RGBColor(0x1A, 0x1A, 0x1A)
    GRAY = RGBColor(0x66, 0x66, 0x66)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(11)
    font.color.rgb = DARK

    # Parse markdown to HTML
    with open(PRD_FILE, 'r', encoding='utf-8') as f:
        md_content = f.read()

    html = markdown.markdown(md_content, extensions=['tables', 'toc', 'fenced_code'])

    # Process HTML line by line
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(24)
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(24)
            run.font.color.rgb = PRIMARY
            run.font.bold = True
            i += 1
            continue

        if line.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.space_before = Pt(18)
            run = p.add_run(line[3:].strip())
            run.font.size = Pt(18)
            run.font.color.rgb = PRIMARY
            run.font.bold = True
            i += 1
            continue

        if line.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(line[4:].strip())
            run.font.size = Pt(14)
            run.font.color.rgb = SECONDARY
            run.font.bold = True
            i += 1
            continue

        if line.startswith('#### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(line[5:].strip())
            run.font.size = Pt(12)
            run.font.color.rgb = SECONDARY
            run.font.bold = True
            i += 1
            continue

        # Horizontal rules
        if line == '---':
            p = doc.add_paragraph()
            p_run = p.add_run('_' * 80)
            p_run.font.color.rgb = GRAY
            i += 1
            continue

        # Tables
        if line.startswith('| '):
            # Gather all table lines
            table_lines = []
            while i < len(lines) and lines[i].strip() and lines[i].strip().startswith('| '):
                table_lines.append(lines[i])
                i += 1

            # Parse table
            if len(table_lines) >= 2:
                # Skip the separator line (second line with dashes)
                header_line = table_lines[0]
                data_lines = [l for l in table_lines[1:] if '---' not in l]

                headers = [cell.strip() for cell in header_line[1:-1].split('|')] if table_lines else []

                # Create table
                num_cols = max(len(headers), max((len(line[1:-1].split('|')) for line in data_lines), default=0))
                table = doc.add_table(rows=1+len(data_lines), cols=num_cols)
                table.style = 'Table Grid'
                table.autofit = True

                # Fill header
                for j, h in enumerate(headers):
                    cell = table.cell(0, j)
                    cell.text = h
                    # Style header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Fill data
                for row_idx, line in enumerate(data_lines):
                    cells = [c.strip() for c in line[1:-1].split('|')]
                    for j, cell_text in enumerate(cells):
                        if j < num_cols:
                            cell = table.cell(row_idx + 1, j)
                            cell.text = cell_text

            continue  # i was incremented in the while loop above

        # Bullet lists
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:].strip())
            run.font.size = Pt(11)
            run.font.color.rgb = DARK
            i += 1
            continue

        # Numbered lists
        if line[0].isdigit() and '. ' in line:
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(line[line.index('. ')+2:].strip())
            run.font.size = Pt(11)
            run.font.color.rgb = DARK
            i += 1
            continue

        # Blockquotes
        if line.startswith('>'):
            p = doc.add_paragraph(style='Quote')
            run = p.add_run(line[1:].strip())
            run.font.italic = True
            run.font.color.rgb = GRAY
            run.font.size = Pt(11)
            i += 1
            continue

        # Regular text
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
        p.paragraph_format.space_after = Pt(4)
        i += 1

    # Save document
    doc.save(OUTPUT_DOCX)
    print(f"Document Word créé avec succès: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


if __name__ == "__main__":
    if not PRD_FILE.exists():
        print(f"Erreur: Le fichier PRD n'existe pas: {PRD_FILE}")
        sys.exit(1)

    output = create_styled_document()
    print(f"✅ PRD TastifyPFA généré avec succès!")
    print(f"📄 Fichier: {output}")
